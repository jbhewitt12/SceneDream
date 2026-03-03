from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .base import BookContent, BookMetadata
from .plain_text_utils import (
    build_chapters_from_paragraphs,
    compute_file_checksum,
    extract_declared_title,
    generate_slug,
    split_wrapped_paragraphs,
    trim_project_gutenberg_boilerplate,
)


class DocxBookLoader:
    """Load DOCX files into the normalized BookContent structure."""

    PARSER_VERSION = "1.0"

    def load(self, path: Path) -> BookContent:
        """Load a DOCX file and return BookContent."""
        if not path.exists():
            raise FileNotFoundError(f"DOCX document not found: {path}")

        warnings: list[str] = []
        parse_errors: list[str] = []
        checksum = compute_file_checksum(path)

        try:
            document = Document(str(path))
        except PackageNotFoundError as exc:
            raise ValueError(f"DOCX document '{path}' is not a valid package.") from exc
        except Exception as exc:  # pragma: no cover - defensive wrapper
            raise ValueError(f"Failed to parse DOCX document '{path}': {exc}") from exc

        lines = [paragraph.text for paragraph in document.paragraphs]
        paragraphs = split_wrapped_paragraphs(lines)
        paragraphs, boilerplate_warnings = trim_project_gutenberg_boilerplate(paragraphs)
        warnings.extend(boilerplate_warnings)

        if not paragraphs:
            raise ValueError(f"DOCX document '{path}' did not contain readable content.")

        title = extract_declared_title(paragraphs) or path.stem
        chapters = build_chapters_from_paragraphs(
            paragraphs=paragraphs,
            default_title=title,
            source_name_prefix=f"{path.stem}_docx",
        )

        metadata = BookMetadata(
            file_path=path,
            file_checksum=checksum,
            parser_version=self.PARSER_VERSION,
            format="docx",
            warnings=warnings,
            parse_errors=parse_errors,
            source_metadata={
                "docx_paragraph_nodes": len(document.paragraphs),
                "paragraph_count": len(paragraphs),
            },
        )

        return BookContent(
            slug=generate_slug(title),
            title=title,
            chapters=chapters,
            metadata=metadata,
            author=None,
        )
